import os
import re
import collections
import docx2txt
import sys
from docx import Document
from docxcompose.composer import Composer
import classes.globals as g


class HsenDocument(object):
    def __init__(self, source_file):
        self.source_file = source_file
        self.source_file_path = os.path.join(g.app.hsen_source_folder, self.source_file)

        self.filename_out = "Processed " + self.source_file
        self.filepath_out = os.path.join(g.app.processed_hsen_folder, self.filename_out)
        self.text_version_file_path = os.path.join(g.app.text_version_folder, self.source_file.replace("docx", "txt"))
        self.terms_file_path = os.path.join(g.app.terms_folder, self.source_file.replace("docx", "txt"))
        self.weighted_terms_file_path = os.path.join(g.app.weighted_terms_folder, self.source_file.replace("docx", "txt"))

    def open(self):
        self.document = Document(self.filepath_out)

    def save(self):
        self.document.save(self.filepath_out)
        pass

    def merge(self):
        master = Document(g.app.template_file)
        composer = Composer(master)
        doc1 = Document(self.source_file_path)
        composer.append(doc1)
        composer.save(self.filepath_out)

    def process(self):
        print("Processing file '{filename}'".format(filename=self.source_file))
        self.merge()
        self.open()
        styles = self.document.styles
        for para in self.document.paragraphs:
            text = para.text
            if text.startswith("Section"):
                para.style = self.document.styles['Title']
            elif text.startswith("Notes."):
                para.style = self.document.styles['Heading 1']
            elif text.startswith("Chapter"):
                para.style = self.document.styles['Heading 1']
            elif text.startswith("Subheading Explanatory Note."):
                para.style = self.document.styles['Heading 2']
            elif text.startswith("Subheading ") or text.startswith("Subheadings "):
                para.style = self.document.styles['Heading 3']
            elif re.match("^[0-9]{2}.[0-9]{2} ", text):
                para.style = self.document.styles['Heading 3']

    def convert_document_to_text_only(self, nlp):
        print("Extracting terms from file '{filename}'".format(filename=self.source_file))
        self.text = docx2txt.process(self.source_file_path)
        doc = nlp(self.text)
        terms = []
        term_dict = {}
        for token in doc:
            if token.pos_ in ("NOUN", "PROPN"):
                value = token.lemma_.lower().replace(".", "")
                if not value.isnumeric():
                    if len(value) > 1:
                        if value not in terms:
                            terms.append(value)

                        if value not in term_dict:
                            term_dict[value] = 1
                        else:
                            term_dict[value] += 1

        # Write the text-only version
        f = open(self.text_version_file_path, "w")
        f.write(self.text)
        f.close()

        # Write the list of terms
        f = open(self.terms_file_path, "w")
        terms.sort()
        for term in terms:
            f.write(term + "\n")
        f.close()

        # Write the weighted list of terms
        term_dict_sorted = collections.OrderedDict(sorted(term_dict.items()))
        with open(self.weighted_terms_file_path, 'w') as f:
            for term in term_dict_sorted:
                line = '"{term}",{count}\n'.format(term=term, count=term_dict_sorted[term])
                f.write(line)
        f.close()

        return terms, term_dict
